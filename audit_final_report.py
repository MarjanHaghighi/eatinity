import sys
import zipfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent / ".final_docx_deps"))
from docx import Document


PATH = Path(__file__).parent / "Eatinity_Final_Project_Report.docx"
doc = Document(PATH)
text = "\n".join(paragraph.text for paragraph in doc.paragraphs)

required = [
    "Solution Overview and Objectives",
    "Solution Design",
    "Architecture Design",
    "Implementation Steps",
    "Relevant Code and Configuration",
    "End-to-End User Experience Walkthroughs",
    "Disaster Recovery Design and Implementation",
    "Automated Security and Code-Quality Testing",
    "GitHub Repository and CI/CD Security",
    "Testing Results",
    "Evidence and Submission Checklist",
]
missing = [heading for heading in required if heading not in text]

with zipfile.ZipFile(PATH) as archive:
    xml = "\n".join(
        archive.read(name).decode("utf-8", errors="ignore")
        for name in archive.namelist()
        if name.endswith(".xml")
    )

secret_markers = ["AKIA", "sk_live_", "sk_test_", "whsec_", "aws_secret_access_key"]
found_secrets = [marker for marker in secret_markers if marker in xml]
placeholder_markers = ["INSERT SCREENSHOT", "REPLACE_WITH", "TODO"]
found_placeholders = [marker for marker in placeholder_markers if marker in text]

section = doc.sections[0]
geometry_ok = all(
    abs(value.inches - expected) < 0.01
    for value, expected in (
        (section.top_margin, 1.0),
        (section.right_margin, 1.0),
        (section.bottom_margin, 1.0),
        (section.left_margin, 1.0),
    )
)

print(f"File: {PATH}")
print(f"Size: {PATH.stat().st_size} bytes")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Headings: {sum(1 for p in doc.paragraphs if p.style.name.startswith('Heading'))}")
print(f"Tables: {len(doc.tables)}")
print(f"Required sections missing: {missing}")
print(f"Secret markers found: {found_secrets}")
print(f"Placeholder markers found: {found_placeholders}")
print(f"Page geometry valid: {geometry_ok}")

if missing or found_secrets or found_placeholders or not geometry_ok:
    raise SystemExit(1)
print("Final report structural QA passed.")

