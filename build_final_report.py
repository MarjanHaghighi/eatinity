import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent / ".final_docx_deps"))

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parent
SOURCE = ROOT / "docs" / "EATINITY_FINAL_PROJECT_REPORT.md"
OUTPUT = ROOT / "Eatinity_Final_Project_Report.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "203748"
GRAY = "555555"
LIGHT = "F2F4F7"
GOLD = "B07D18"


def set_font(run, size=11, bold=False, italic=False, color="000000"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_pr = table._tbl.tblPr
    tbl_w = table_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        table_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = table_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        table_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def add_cover(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("EATINITY | FINAL PROJECT"), size=9, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(120)
    p.paragraph_format.space_after = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("FINAL PROJECT TECHNICAL REPORT"), size=11, bold=True, color=GOLD)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Eatinity Secure E-Commerce Platform"), size=28, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("Serverless AWS commerce, secure CI/CD, and cross-region recovery"), size=14, color=DARK_BLUE)

    for label, value in (
        ("Application", "https://eatinity.ca"),
        ("Platform", "Amazon Web Services"),
        ("Prepared for", "David - Final Project Submission"),
        ("Date", "August 1, 2026"),
    ):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        set_font(p.add_run(f"{label}: "), size=10.5, bold=True, color=GRAY)
        set_font(p.add_run(value), size=10.5, color=GRAY)
    doc.add_page_break()


def add_inline(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_font(run, size=9.5, color=DARK_BLUE)
            run.font.name = "Consolas"
            run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
        elif part.startswith("**") and part.endswith("**"):
            set_font(paragraph.add_run(part[2:-2]), bold=True)
        else:
            set_font(paragraph.add_run(part))


def add_markdown(doc, lines):
    i = 0
    in_code = False
    code_lines = []
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            if in_code:
                table = doc.add_table(rows=1, cols=1)
                table.style = "Table Grid"
                set_table_geometry(table, [9360])
                shade(table.cell(0, 0), LIGHT)
                p = table.cell(0, 0).paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                run = p.add_run("\n".join(code_lines))
                set_font(run, size=8.5, color=NAVY)
                run.font.name = "Consolas"
                run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Consolas")
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[-| :]+\|$", lines[i + 1].strip()):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [cell.strip() for cell in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r"[-: ]+", cell) for cell in cells):
                    rows.append(cells)
                i += 1
            column_count = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=column_count)
            table.style = "Table Grid"
            widths = [9360 // column_count] * column_count
            widths[-1] += 9360 - sum(widths)
            for r_index, row in enumerate(rows):
                for c_index, value in enumerate(row):
                    cell = table.cell(r_index, c_index)
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline(p, value)
                    if r_index == 0:
                        shade(cell, LIGHT)
                        for run in p.runs:
                            run.bold = True
            set_table_geometry(table, widths)
            continue
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:], level=3)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, line[2:])
        elif re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\. ", "", line))
        else:
            p = doc.add_paragraph()
            add_inline(p, line.replace("  ", " "))
        i += 1


def main():
    doc = Document()
    configure_styles(doc)
    add_cover(doc)
    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    add_markdown(doc, lines)
    props = doc.core_properties
    props.title = "Eatinity Secure E-Commerce Platform - Final Project Technical Report"
    props.subject = "AWS solution, implementation, CI/CD, security, user workflows, and disaster recovery"
    props.author = "Eatinity Project Group"
    props.keywords = "AWS, Eatinity, GitHub Actions, OIDC, Trivy, SonarQube, Disaster Recovery"
    doc.save(OUTPUT)
    print(f"Created {OUTPUT}")


if __name__ == "__main__":
    main()
