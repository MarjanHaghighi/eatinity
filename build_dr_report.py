import sys
sys.path.insert(0, r"C:\Marjan\Eatinity\.docx_deps")

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from pathlib import Path

OUT = Path(r"C:\Marjan\Eatinity\Eatinity_Disaster_Recovery_Plan.docx")
NAVY = "17365D"
BLUE = "2E74B5"
TEAL = "1F6D73"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "6B7280"
GREEN = "217346"
GOLD = "A66A00"
RED = "9B1C1C"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)


def set_font(run, size=None, bold=None, color=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MID_GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p


def add_callout(doc, label, text, fill=LIGHT_BLUE, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=color)
    r = p.add_run(text)
    set_font(r, color=color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_figure_placeholder(doc, number, title, instruction):
    caption = doc.add_paragraph()
    caption.paragraph_format.space_before = Pt(8)
    caption.paragraph_format.space_after = Pt(4)
    r = caption.add_run(f"Figure {number}. {title}")
    set_font(r, bold=True, size=10, color=NAVY)
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F8FAFC")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(24)
    r = p.add_run(f"INSERT SCREENSHOT {number} HERE\n")
    set_font(r, bold=True, size=13, color=BLUE)
    r = p.add_run(instruction)
    set_font(r, size=9.5, color=MID_GRAY, italic=True)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(6)
    r = note.add_run("Evidence note: ")
    set_font(r, bold=True, size=9, color=MID_GRAY)
    r = note.add_run("Crop sensitive values such as Stripe keys, webhook secrets, and temporary passwords before insertion.")
    set_font(r, size=9, color=MID_GRAY)


def add_table(doc, headers, rows, widths, font_size=9.5):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    set_repeat_table_header(header)
    for i, text in enumerate(headers):
        cell = header.cells[i]
        set_cell_shading(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(text)
        set_font(r, bold=True, color="FFFFFF", size=font_size)
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if len(table.rows) % 2 == 1:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_font(r, size=font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 11.5, NAVY, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for style_name in ("List Bullet", "List Bullet 2", "List Number"):
    s = styles[style_name]
    s.font.name = "Calibri"
    s.font.size = Pt(10.5)
    s.paragraph_format.space_after = Pt(5)
    s.paragraph_format.line_spacing = 1.10

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = hp.add_run("EATINITY | DISASTER RECOVERY PLAN")
set_font(r, bold=True, size=9, color=TEAL)
footer = section.footer
fp = footer.paragraphs[0]
add_page_number(fp)

# Cover: editorial report pattern.
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(68)
p.paragraph_format.space_after = Pt(12)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("FINAL PROJECT DELIVERY")
set_font(r, bold=True, size=10, color=TEAL)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(10)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Eatinity Disaster Recovery Plan")
set_font(r, bold=True, size=28, color=NAVY)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(28)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Cross-region recovery automation for AWS infrastructure, data, identity, email, and application services")
set_font(r, size=13, color=MID_GRAY)

add_callout(doc, "Recovery drill outcome", "The Eatinity workload was rebuilt and validated in ca-central-1 from a successful us-east-1 source environment. Terraform reported zero destroys, seven AWS Backup recovery points were copied cross-region, five DynamoDB datasets were restored and validated, Cognito users were recovered, SES domain/DKIM verification succeeded, and the regional application completed functional testing.", fill="E9F5EE", color=GREEN)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(30)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, value in (
    ("Prepared for", "David - Final Project / Disaster Recovery submission"),
    ("Application", "Eatinity"),
    ("Source region", "US East (N. Virginia) - us-east-1"),
    ("Recovery region", "Canada (Central) - ca-central-1"),
    ("Recovery drill date", "July 21-22, 2026"),
):
    run = p.add_run(f"{label}: ")
    set_font(run, bold=True, size=10.5, color=NAVY)
    run = p.add_run(value + "\n")
    set_font(run, size=10.5)

doc.add_page_break()

doc.add_heading("1. Executive Summary", level=1)
doc.add_paragraph("This Disaster Recovery Plan explains how Eatinity and its supporting AWS infrastructure can be restored after a regional failure or major outage. The plan uses Terraform for repeatable infrastructure provisioning and AWS Backup for native, cross-region protection of Amazon DynamoDB and Amazon S3. Recovery automation is implemented through reviewed PowerShell runbooks that test prerequisites, start backups, copy recovery points, restore into isolated drill resources, validate record counts, and promote restored data into the regional application tables.")
doc.add_paragraph("The recovery drill preserved the original us-east-1 environment. All recovery infrastructure used a region-derived naming prefix, eatinity-prod-cac1, and was created in ca-central-1. The drill also recovered Cognito users and groups, established a controlled password-reset process, configured regional SES identity and DKIM, and deployed the frontend using dynamic Terraform outputs rather than hard-coded endpoints.")

doc.add_heading("1.1 Requirement Coverage", level=2)
add_table(doc, ["Instructor requirement", "Implemented evidence", "Status"], [
    ("Infrastructure Recovery", "Terraform modules, regional variables, saved plans, apply output, and repository screenshot placeholder", "Complete"),
    ("Database Recovery", "AWS Backup plan/vaults, cross-region copies, five DynamoDB restores, S3 restores, and record-count validation", "Complete"),
    ("Application Validation", "Regional CloudFront deployment, Cognito sign-in recovery, product/order/Stripe workflow tests, SES verification", "Complete"),
    ("Security Scanning", "GitHub Actions workflow using Trivy for vulnerabilities, secrets, dependencies, IaC misconfiguration, and optional images", "Implemented; push to GitHub to execute"),
], [2500, 5360, 1500])

doc.add_heading("1.2 Recovery Objectives", level=2)
add_table(doc, ["Objective", "Target / implementation", "Evidence"], [
    ("Recovery Point Objective (RPO)", "24 hours maximum under the scheduled daily backup rule; lower for on-demand drill backups", "AWS Backup cron schedule at 05:00 UTC"),
    ("Recovery Time Objective (RTO)", "Target 4 hours for an operator-led regional restoration", "Drill demonstrated parallel backup copy, Terraform deployment, restore, and validation"),
    ("Retention", "35 days in source vault; 90 days in recovery-region vault", "Terraform recovery_backup module"),
    ("Isolation", "Restore into temporary drill resources before controlled promotion", "eatinity-recovery-drill-* tables and validation scripts"),
], [1800, 4860, 2700])

doc.add_heading("2. Recovery Architecture", level=1)
doc.add_paragraph("The recovery design separates infrastructure configuration, protected data, restored data, and validation evidence. Terraform creates the recovery services and permissions. AWS Backup stores native recovery points in the source vault and copies them to a destination vault in ca-central-1. Restores first land in isolated drill resources so their structure and contents can be checked without overwriting active destination resources.")
add_figure_placeholder(doc, 1, "Cross-region disaster recovery architecture", "Insert an AWS architecture diagram or console screenshot showing us-east-1 source resources, AWS Backup source vault, cross-region copy, ca-central-1 recovery vault, restored DynamoDB/S3, Cognito, SES, API Gateway, Lambda, CloudFront, and the regional frontend.")

doc.add_heading("2.1 Protected and Rebuilt Services", level=2)
add_table(doc, ["Layer", "Source / protection", "Recovery implementation"], [
    ("Infrastructure", "Terraform configuration and variables", "Recreate 99 AWS resources using regional and account-aware names"),
    ("DynamoDB", "Five source tables protected by AWS Backup", "Restore drill tables, validate, then copy 97 records to five destination tables"),
    ("S3", "Website and images buckets with versioning", "Native S3 backup, cross-region recovery points, and restore to regional buckets"),
    ("Identity", "Cognito user pool in us-east-1", "Synchronize users, attributes, status, groups; controlled password re-establishment"),
    ("Email", "SES domain identity and templates", "Regional SES identity, DKIM, MAIL FROM, SNS bounce/complaint monitoring"),
    ("Application", "React frontend, API Gateway, Lambda, Stripe test integration", "Dynamic output-driven deployment to regional S3/CloudFront"),
], [1500, 3500, 4360])

doc.add_page_break()
doc.add_heading("3. Infrastructure Recovery with IaC", level=1)
doc.add_paragraph("Terraform is the authoritative recovery mechanism for the AWS platform. The production environment configuration is reusable across regions because the AWS region, project name, environment name, account ID, resource prefix, URLs, bucket names, Cognito identifiers, and frontend endpoints are calculated from variables, data sources, and module outputs. Region-specific strings are not embedded throughout the resource files.")

doc.add_heading("3.1 Repository Organization", level=2)
add_bullet(doc, "environments/production - regional variables, providers, checks, outputs, and module composition")
add_bullet(doc, "modules/application - API Gateway, Lambda, IAM, CloudWatch, and application integrations")
add_bullet(doc, "modules/database - DynamoDB tables, indexes, PITR, and deletion protection")
add_bullet(doc, "modules/storage - S3 website/images buckets, versioning, encryption, and CORS")
add_bullet(doc, "modules/identity - Cognito user pool, client, and authorizer configuration")
add_bullet(doc, "modules/delivery - CloudFront distribution and optional public DNS")
add_bullet(doc, "modules/operations - SES, DKIM, MAIL FROM, SNS, and notification configuration")
add_bullet(doc, "modules/recovery_backup - AWS Backup vaults, plan, cross-region copy, IAM role, and operator policy")
add_bullet(doc, "migration - prerequisite, backup, copy, restore, validation, Cognito, SES, and deployment scripts")
add_figure_placeholder(doc, 2, "IaC repository structure", "Insert a VS Code Explorer screenshot showing eatinity-iac/environments/production, modules, migration, README files, and the GitHub Actions security workflow. Ensure secrets and terraform.tfvars values are not visible.")

doc.add_heading("3.2 Terraform Deployment Procedure", level=2)
for step in [
    "Select the recovery environment and review terraform.tfvars. The tested destination is ca-central-1 and the source region is us-east-1.",
    "Run terraform validate to verify syntax and provider/module references.",
    "Create a saved plan with terraform plan -out=<reviewed-name>.tfplan.",
    "Review the summary and confirm that planned destruction is zero.",
    "Apply the exact saved plan with terraform apply <reviewed-name>.tfplan.",
    "Capture Terraform outputs for API, CloudFront, Cognito, buckets, backup vaults, and SES.",
    "Run a final terraform plan and confirm that configuration drift is absent.",
]:
    add_number(doc, step)

doc.add_heading("3.3 Deployment Results", level=2)
add_callout(doc, "Initial regional deployment", "Terraform planned and created 99 resources in ca-central-1 with 0 changes and 0 destroys. Later reviewed plans added AWS Backup, operator permissions, dynamic frontend outputs, Cognito recovery permissions, and SES recovery resources. The final SES plan reported 16 additions, 1 IAM update, and 0 destroys.", fill=LIGHT_BLUE)
add_figure_placeholder(doc, 3, "Terraform plan and apply evidence", "Insert the terminal screenshot showing a reviewed plan/apply result. Recommended evidence: 'Plan: 99 to add, 0 to change, 0 to destroy' plus the output block, or the final 'No changes' plan.")
add_figure_placeholder(doc, 4, "Regional AWS resources", "Insert an AWS Console screenshot from Canada (Central) showing the Eatinity resources and region selector. Good choices: CloudFormation-style resource list, DynamoDB tables, S3 buckets, API Gateway, or CloudFront distribution.")

doc.add_heading("3.4 Safety Controls", level=2)
add_bullet(doc, "A recovery-region guard blocks accidental deployment to us-east-1 unless explicitly authorized.")
add_bullet(doc, "Bucket force deletion is disabled and DynamoDB deletion protection is enabled.")
add_bullet(doc, "Source data operations are read-only except for starting AWS Backup jobs.")
add_bullet(doc, "Restore operations target isolated resources before promotion.")
add_bullet(doc, "IAM PassRole is restricted to the dedicated AWS Backup role and backup.amazonaws.com.")
add_bullet(doc, "Cognito source permissions are read-only; destination permissions are limited to recovery user/group operations.")
add_bullet(doc, "terraform.tfvars and state/plan artifacts are excluded from version control because they may contain sensitive values.")

doc.add_page_break()
doc.add_heading("4. Database and Object Storage Recovery", level=1)
doc.add_heading("4.1 Backup Strategy", level=2)
doc.add_paragraph("AWS Backup provides a native, policy-driven backup strategy for five DynamoDB tables and two S3 buckets. The plan runs daily at 05:00 UTC, retains source-region recovery points for 35 days, and copies recovery points to the ca-central-1 vault for 90 days. S3 versioning is a prerequisite. DynamoDB point-in-time recovery remains enabled as an additional table-level protection mechanism.")

add_table(doc, ["Protected resource", "Source name", "Recovery destination"], [
    ("DynamoDB audit", "EatinityAuditLog", "eatinity-prod-cac1-audit"),
    ("DynamoDB categories", "EatinityCategories", "eatinity-prod-cac1-categories"),
    ("DynamoDB orders", "EatinityOrders", "eatinity-prod-cac1-orders"),
    ("DynamoDB products", "EatinityProducts", "eatinity-prod-cac1-products"),
    ("DynamoDB users", "EatinityUsers", "eatinity-prod-cac1-users"),
    ("S3 images", "eatinity-prod-s3-images", "eatinity-prod-cac1-969433238478-images"),
    ("S3 website", "eatinity-prod-s3-website", "eatinity-prod-cac1-969433238478-website"),
], [1900, 3400, 4060])

doc.add_heading("4.2 Automated Backup and Cross-Region Copy", level=2)
for step in [
    "Test-NativeBackupPrerequisites.ps1 verifies all source tables, S3 versioning, and AWS Backup service opt-in.",
    "Start-NativeBackup.ps1 starts one native backup job per protected resource and records job metadata.",
    "Get-NativeRecoveryJobStatus.ps1 monitors backup jobs until completion.",
    "Start-NativeBackupCopy.ps1 copies each recovery point to eatinity-prod-cac1-recovery-vault in ca-central-1.",
    "Copy status is monitored until all seven jobs are completed.",
]:
    add_number(doc, step)
add_figure_placeholder(doc, 5, "Seven completed cross-region recovery points", "Insert the AWS CLI or AWS Backup console screenshot showing all five DynamoDB and two S3 copy jobs as COMPLETED in ca-central-1.")

doc.add_heading("4.3 DynamoDB Restore and Promotion", level=2)
doc.add_paragraph("Each recovery point was restored first into an isolated table named eatinity-recovery-drill-<logical-name>-20260721. The restored tables were inspected and compared with the source. Copy-RestoredDynamoData.ps1 then promoted the validated records into the Terraform-managed destination tables. The script is idempotent: tables with the expected count are skipped, and evidence is written to migration/job-records.")
add_table(doc, ["Dataset", "Restored status", "Records validated", "Destination result"], [
    ("Audit", "ACTIVE", "32", "PASS"),
    ("Categories", "ACTIVE", "8", "PASS"),
    ("Orders", "ACTIVE", "34", "PASS; 2 empty secondary-index keys sanitized"),
    ("Products", "ACTIVE", "21", "PASS"),
    ("Users", "ACTIVE", "2", "PASS"),
], [1900, 1800, 2100, 3560])
add_figure_placeholder(doc, 6, "DynamoDB restored drill tables", "Insert the DynamoDB Tables screenshot listing all five eatinity-recovery-drill-* tables in ACTIVE state.")
add_figure_placeholder(doc, 7, "DynamoDB record-count validation", "Insert the PowerShell output showing Audit 32/32, Categories 8/8, Orders 34/34, Products 21/21, Users 2/2, followed by the successful destination copy evidence.")

doc.add_heading("4.4 S3 Restore", level=2)
doc.add_paragraph("The website and image buckets were backed up using AWS Backup for S3, copied into the ca-central-1 vault, and restored to regional S3 resources. Versioning was enabled before backup. Recovery jobs were monitored by job ID and recovery-point ARN until they completed. This native approach scales better than exporting object lists to local JSON files and keeps recovery metadata within AWS Backup.")
add_figure_placeholder(doc, 8, "S3 backup and restore completion", "Insert the AWS Backup console or CLI screenshot showing the website and image S3 jobs as COMPLETED, with resource type S3 and destination region ca-central-1.")

doc.add_page_break()
doc.add_heading("5. Identity and Email Recovery", level=1)
doc.add_heading("5.1 Cognito Recovery", level=2)
doc.add_paragraph("Cognito passwords and MFA secrets cannot be exported. The recovery automation therefore transfers only recoverable identity data: users, standard attributes, enabled/disabled status, groups, and group membership. The source user pool is read-only. Users are created in the destination pool with messages suppressed during bulk synchronization; a controlled invitation/reset process is then used to establish a new password.")
add_table(doc, ["Validation item", "Result"], [
    ("Source users", "2"),
    ("Destination users", "2"),
    ("Missing source users", "0"),
    ("Source group memberships", "1"),
    ("Destination group memberships", "1"),
    ("Validation", "PASS"),
], [4300, 5060])
add_figure_placeholder(doc, 9, "Cognito synchronization and validation", "Insert the Test-CognitoRecoveryUsers.ps1 screenshot showing 2 source users, 2 destination users, 0 missing users, matching group memberships, and Validation PASS.")

doc.add_heading("5.2 Cognito Password Challenge", level=2)
doc.add_paragraph("The React frontend was enhanced to support Cognito's NEW_PASSWORD_REQUIRED challenge. A recovered user signs in with the temporary invitation password, enters and confirms a permanent password, and then receives a valid regional Cognito session. The frontend does not hard-code passwords or store them in Terraform.")
add_figure_placeholder(doc, 10, "Recovered user password and sign-in test", "Insert the 'Choose New Password' page and/or the successfully authenticated recovery application page. Do not show any password or temporary credential.")

doc.add_heading("5.3 SES Regional Recovery", level=2)
doc.add_paragraph("SES was configured independently in ca-central-1 because identities, sending quotas, and production-access status are regional. Terraform created the eatinity.ca domain identity, three DKIM records, a regional custom MAIL FROM domain, a configuration set, and separate SNS topics/subscriptions for bounce and complaint events. Template synchronization is automated; the source and destination template counts are both zero because the application currently sends non-template messages.")
add_table(doc, ["SES validation", "Result"], [
    ("Region", "ca-central-1"),
    ("Domain verification", "Success"),
    ("DKIM enabled / verified", "True / Success"),
    ("Sending enabled", "True"),
    ("Production access", "False - regional SES sandbox limitation"),
    ("Template count", "0 source / 0 destination"),
], [4300, 5060])
add_figure_placeholder(doc, 11, "SES domain and DKIM validation", "Insert the Test-SesRecovery.ps1 screenshot showing VerificationStatus Success, DkimVerificationStatus Success, and SendingEnabled True.")
add_callout(doc, "Known limitation", "ProductionAccessEnabled is False because ca-central-1 remains in the SES sandbox. This does not invalidate infrastructure recovery. A real production cutover requires an approved SES production-access request and verified sending/receiving identities while sandboxed.", fill="FFF4E5", color=GOLD)

doc.add_page_break()
doc.add_heading("6. Application Deployment and Validation", level=1)
doc.add_heading("6.1 Dynamic Regional Deployment", level=2)
doc.add_paragraph("Deploy-RegionalFrontend.ps1 builds the React application, reads Terraform outputs, writes a runtime configuration containing the regional API endpoint, image bucket, AWS region, Cognito pool ID, and client ID, uploads the build to the destination S3 website bucket, and creates a CloudFront invalidation. This makes the same application package reusable in another region without source-code endpoint changes.")

doc.add_heading("6.2 Functional Validation Results", level=2)
add_table(doc, ["Test", "Expected result", "Observed result"], [
    ("CloudFront availability", "Regional URL loads over HTTPS", "PASS"),
    ("Products/categories", "Restored catalogue is displayed", "PASS"),
    ("Cognito sign-in", "Recovered user can establish a new password and authenticate", "PASS"),
    ("Cart/checkout", "Items can be selected and checkout starts", "PASS"),
    ("Stripe test payment", "Stripe test session completes and redirects to regional success page", "PASS"),
    ("Order data", "Order is written/read from ca-central-1 DynamoDB", "PASS"),
    ("Images/static assets", "Assets load from regional S3/CloudFront", "PASS"),
    ("SES identity", "Regional identity and DKIM validate", "PASS; sandbox remains"),
], [2100, 4200, 3060])
add_figure_placeholder(doc, 12, "Restored regional application", "Insert the ca-central-1 CloudFront application home/products page. Keep the CloudFront URL visible in the browser address bar.")
add_figure_placeholder(doc, 13, "Stripe checkout and success validation", "Insert the successful Stripe test payment or Thank You page from the recovery CloudFront URL. Ensure no payment card details are visible.")
add_figure_placeholder(doc, 14, "Final IaC drift check", "Insert the final terraform plan output showing 'No changes. Your infrastructure matches the configuration.' If not yet captured, run it only after all intended changes are applied.")

doc.add_heading("7. Security Scanning in CI/CD", level=1)
doc.add_paragraph("A GitHub Actions workflow has been added at .github/workflows/security-scan.yml. It runs on pushes and pull requests to main and can also be started manually. The required filesystem job checks the codebase, dependency lockfiles, embedded secrets, and Terraform/AWS configuration for HIGH and CRITICAL findings. A manually supplied container image can be scanned by the optional image job. A qualifying finding returns a non-zero exit code and blocks the workflow before deployment.")
add_table(doc, ["Scan coverage", "Trivy mode", "Pipeline behavior"], [
    ("Codebase and dependencies", "Filesystem vulnerability scan", "Fail on HIGH/CRITICAL unfixed findings excluded"),
    ("Secrets", "Filesystem secret scanner", "Detect credentials accidentally committed to Git"),
    ("Terraform/IaC", "Misconfiguration scanner", "Identify insecure AWS/Terraform settings before deployment"),
    ("Container image", "Optional workflow_dispatch image scan", "Scan supplied image reference before release"),
], [2300, 3200, 3860])
add_figure_placeholder(doc, 15, "GitHub Actions Trivy security scan", "After committing and pushing, open GitHub > Actions > Security Scan. Insert the workflow summary showing the Trivy source/dependency/secret/IaC job and its result.")
add_callout(doc, "Required follow-up", "Commit and push the workflow to GitHub, correct any HIGH/CRITICAL findings, rerun it to green, and capture Figure 15. The local file alone does not prove that the hosted CI job executed.", fill="FFF4E5", color=GOLD)

doc.add_page_break()
doc.add_heading("8. End-to-End Recovery Runbook", level=1)
for step in [
    "Declare the incident, identify the unavailable region/services, and authorize recovery in ca-central-1.",
    "Confirm AWS credentials, account ID, source region, destination region, and recovery configuration.",
    "Run Terraform validate, plan, review, and apply to rebuild the regional infrastructure.",
    "Run Test-NativeBackupPrerequisites.ps1 and select the latest valid recovery points.",
    "Copy required recovery points to eatinity-prod-cac1-recovery-vault if not already present.",
    "Restore DynamoDB and S3 into isolated drill resources; monitor job status to completion.",
    "Validate DynamoDB schemas, indexes, counts, and sample records; validate S3 object availability.",
    "Promote validated data into the Terraform-managed destination tables and buckets.",
    "Synchronize Cognito users/groups, validate identity counts, and initiate controlled password establishment.",
    "Apply and validate SES domain/DKIM settings; confirm alert subscriptions and regional sending eligibility.",
    "Deploy the frontend with Deploy-RegionalFrontend.ps1 using dynamic Terraform outputs.",
    "Test authentication, catalogue, images, checkout, Stripe test payment, order persistence, logs, and email behavior.",
    "Run a final Terraform drift check, archive job records/screenshots, and obtain stakeholder approval for DNS cutover if required.",
]:
    add_number(doc, step)

doc.add_heading("9. Evidence Register and Screenshot Checklist", level=1)
add_table(doc, ["Fig.", "Required screenshot", "Suggested source", "Added?"], [
    ("1", "Cross-region architecture", "AWS diagram / draw.io", "[ ]"),
    ("2", "IaC repository", "VS Code Explorer", "[ ]"),
    ("3", "Terraform plan/apply", "PowerShell terminal", "[ ]"),
    ("4", "Regional resources", "AWS Console ca-central-1", "[ ]"),
    ("5", "Seven completed copies", "AWS Backup / CLI", "[ ]"),
    ("6", "Five drill tables", "DynamoDB Console", "[ ]"),
    ("7", "Record validation", "PowerShell output", "[ ]"),
    ("8", "S3 restore", "AWS Backup / S3", "[ ]"),
    ("9", "Cognito PASS", "PowerShell output", "[ ]"),
    ("10", "Password/sign-in test", "Recovery frontend", "[ ]"),
    ("11", "SES/DKIM success", "PowerShell output", "[ ]"),
    ("12", "Regional app", "Browser", "[ ]"),
    ("13", "Stripe success", "Browser", "[ ]"),
    ("14", "No-change plan", "PowerShell terminal", "[ ]"),
    ("15", "Trivy CI result", "GitHub Actions", "[ ]"),
], [600, 3000, 4100, 1660], font_size=8.7)

doc.add_heading("10. Automation and Evidence Files", level=1)
doc.add_paragraph("The following files form the reproducible evidence package and should be submitted with or referenced by the report:")
add_bullet(doc, "eatinity-iac/environments/production - Terraform environment, saved plans, variables, checks, and outputs")
add_bullet(doc, "eatinity-iac/modules/recovery_backup - backup vaults, daily plan, cross-region copy, IAM role, and least-privilege operator policy")
add_bullet(doc, "eatinity-iac/modules/operations - SES identity, DKIM, MAIL FROM, SNS bounce/complaint monitoring")
add_bullet(doc, "eatinity-iac/migration - automated prerequisite, backup, copy, restore, validation, deployment, Cognito, and SES scripts")
add_bullet(doc, "eatinity-iac/migration/job-records - timestamped JSON evidence for backup, copy, restore, DynamoDB promotion, Cognito sync, and SES template sync")
add_bullet(doc, ".github/workflows/security-scan.yml - Trivy security gate for CI/CD")

doc.add_heading("11. Limitations and Improvement Actions", level=1)
add_bullet(doc, "Cognito passwords and MFA secrets cannot be copied. Users must establish a password in the recovery pool; MFA re-enrolment must be included in a production cutover plan.")
add_bullet(doc, "SES production access is regional and ca-central-1 remains in the sandbox until AWS approves the request.")
add_bullet(doc, "A public DNS failover was intentionally not performed during the drill. The regional CloudFront URL was used to avoid changing the live eatinity.ca application.")
add_bullet(doc, "The current RTO target is an operational objective; future drills should capture start/end timestamps per stage to calculate measured RTO.")
add_bullet(doc, "Trivy must run in GitHub after the workflow is pushed. Findings must be remediated or formally accepted before deployment.")
add_bullet(doc, "Recovery drills should be scheduled quarterly, with evidence retained and the runbook updated after each exercise.")

doc.add_heading("12. Conclusion", level=1)
doc.add_paragraph("The Eatinity recovery drill demonstrated that the application can be rebuilt in a separate AWS region using Infrastructure as Code and restored using native AWS backup services. The source environment remained operational and unchanged. Infrastructure deployment, seven cross-region data copies, five DynamoDB restores, S3 restoration, Cognito identity recovery, SES regional verification, and end-to-end application tests produced repeatable scripts and evidence. The remaining operational items are SES production-access approval, hosted execution of the Trivy workflow, final screenshot insertion, and formal approval before any public DNS cutover.")

doc.core_properties.title = "Eatinity Disaster Recovery Plan"
doc.core_properties.subject = "Infrastructure, database, application, identity, email, and CI/CD recovery"
doc.core_properties.author = "Eatinity Project Team"
doc.core_properties.keywords = "AWS, Terraform, Disaster Recovery, DynamoDB, S3, Cognito, SES, Trivy"
doc.save(OUT)
print(OUT)
