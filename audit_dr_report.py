import sys
import zipfile
from pathlib import Path

sys.path.insert(0, r"C:\Marjan\Eatinity\.docx_deps")
from docx import Document

path = Path(r"C:\Marjan\Eatinity\Eatinity_Disaster_Recovery_Plan.docx")
doc = Document(path)
text = "\n".join(p.text for p in doc.paragraphs)
required = [
    "Infrastructure Recovery with IaC",
    "Database and Object Storage Recovery",
    "Application Deployment and Validation",
    "Security Scanning in CI/CD",
    "End-to-End Recovery Runbook",
    "Evidence Register and Screenshot Checklist",
]
missing = [item for item in required if item not in text]
with zipfile.ZipFile(path) as archive:
    xml_text = "".join(
        archive.read(name).decode("utf-8", errors="ignore")
        for name in archive.namelist()
        if name.endswith(".xml")
    )
secret_markers = ["sk_test_", "whsec_", "REPLACE_WITH_STRIPE"]
found_secrets = [marker for marker in secret_markers if marker in xml_text]
placeholders = sum(1 for table in doc.tables for cell in table._cells if "INSERT SCREENSHOT" in cell.text)
print(f"File: {path}")
print(f"Size: {path.stat().st_size} bytes")
print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")
print(f"Screenshot placeholders: {placeholders}")
print(f"Missing required sections: {missing}")
print(f"Secret markers found: {found_secrets}")
if missing or found_secrets or placeholders != 15:
    raise SystemExit(1)
print("Structural QA passed.")
