from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).parent / ".final_docx_deps"))

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parent
OUT = ROOT / "deliverables"
OUT.mkdir(exist_ok=True)

FONT = "Times New Roman"
INK = "000000"
NAVY = "17365D"
BLUE = "1F4E79"
GRAY = "666666"
LIGHT = "F2F2F2"
PALE_BLUE = "EAF2F8"
PALE_GOLD = "FFF4CC"
GREEN = "E2F0D9"


def font(run, size=11, bold=False, italic=False, color=INK):
    run.font.name = FONT
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), FONT)
    rpr.rFonts.set(qn("w:hAnsi"), FONT)
    rpr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    return run


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(paragraph.add_run("Page "), 9, color=GRAY)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def configure(doc, running_title):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Title", 22, NAVY, 0, 12),
        ("Subtitle", 13, GRAY, 0, 10),
        ("Heading 1", 16, NAVY, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11.5, BLUE, 10, 4),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(3)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    font(header.add_run(running_title), 9, bold=True, color=GRAY)
    add_page_number(section.footer.paragraphs[0])


def p(doc, text="", bold_lead=None, italic=False, align=None, after=6):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(after)
    if align is not None:
        para.alignment = align
    if bold_lead and text.startswith(bold_lead):
        font(para.add_run(bold_lead), bold=True)
        font(para.add_run(text[len(bold_lead):]), italic=italic)
    else:
        font(para.add_run(text), italic=italic)
    return para


def bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    font(para.add_run(text))
    return para


def number(doc, text):
    para = doc.add_paragraph(style="List Number")
    font(para.add_run(text))
    return para


def heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_pr = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_pr.append(repeat)
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT)
        para = cell.paragraphs[0]
        para.paragraph_format.line_spacing = 1.15
        font(para.add_run(value), 10, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            para = cells[i].paragraphs[0]
            para.paragraph_format.line_spacing = 1.15
            font(para.add_run(str(value)), 10)
    if widths is None:
        widths = [9360 // len(headers)] * len(headers)
        widths[-1] += 9360 - sum(widths)
    table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def numbered_steps(doc, items):
    for index, text in enumerate(items, start=1):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)
        para.paragraph_format.line_spacing = 1.5
        font(para.add_run(f"Step {index}. "), bold=True)
        font(para.add_run(text))


def evidence(doc, evidence_id, title, capture, purpose):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, PALE_GOLD)
    para = cell.paragraphs[0]
    para.paragraph_format.line_spacing = 1.15
    para.paragraph_format.space_after = Pt(3)
    font(para.add_run(f"SCREENSHOT PLACEHOLDER {evidence_id}: {title}\n"), 10, bold=True, color=NAVY)
    font(para.add_run(f"Add: {capture}\n"), 9.5)
    font(para.add_run(f"Why: {purpose}"), 9.5, italic=True, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def show_box(doc, label, text, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, fill)
    para = cell.paragraphs[0]
    para.paragraph_format.line_spacing = 1.15
    font(para.add_run(f"{label}: "), 10, bold=True, color=NAVY)
    font(para.add_run(text), 10)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def cover(doc, title, subtitle):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(92)
    spacer.paragraph_format.space_after = Pt(8)
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(para.add_run("CAA900 FINAL PROJECT"), 11, bold=True, color=BLUE)
    para = doc.add_paragraph(style="Title")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(para.add_run(title), 22, bold=True, color=NAVY)
    para = doc.add_paragraph(style="Subtitle")
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(para.add_run(subtitle), 13, color=GRAY)
    doc.add_paragraph()
    for label, value in (
        ("Project", "eatinity Secure E-Commerce Platform"),
        ("Student/Group", "[Insert student names and group number]"),
        ("Professor", "[Insert professor name]"),
        ("Course", "CAA900 - Cloud Architecture and Administration"),
        ("Repository", "https://github.com/MarjanHaghighi/eatinity"),
        ("Application", "https://eatinity.ca"),
        ("Submission date", "[Insert submission date]"),
    ):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(3)
        font(para.add_run(f"{label}: "), 10.5, bold=True)
        font(para.add_run(value), 10.5)
    doc.add_page_break()


def report_document():
    doc = Document()
    configure(doc, "eatinity Secure E-Commerce Platform | Final Report")
    cover(doc, "Final Project Technical Report", "Architecture, implementation, security, CI/CD, operations, and disaster recovery")

    heading(doc, "Document Completion Instructions", 1)
    p(doc, "This document is written as a submission-ready report. Replace every item in square brackets, insert the requested screenshots in the labeled locations, update the table of contents in Microsoft Word, and remove this instruction section before final submission. Never insert passwords, access keys, secret values, tokens, private keys, full payment details, or personally identifiable customer data.")
    show_box(doc, "Evidence rule", "Each major claim should be explained, demonstrated, and supported by readable evidence. Crop screenshots to the relevant result, keep the service name and status visible, and add a one-sentence caption with the date and result.", GREEN)

    heading(doc, "Table of Contents", 1)
    for item in (
        "1. Solution Overview", "2. Solution Design", "3. Architecture", "4. Implementation", "5. Important Code and Configuration", "6. User Walkthroughs", "7. Disaster Recovery", "8. Automated Security", "9. GitHub and AWS OIDC", "10. AWS Well-Architected Review", "11. Testing and Results", "12. Risks, Limitations, and Future Improvements", "13. Conclusion", "Appendix A. Evidence Checklist", "Appendix B. Repository Structure and References"
    ):
        p(doc, item, after=1)
    doc.add_page_break()

    heading(doc, "1. Solution Overview", 1)
    heading(doc, "1.1 Problem and purpose", 2)
    p(doc, "eatinity is a secure, serverless e-commerce platform for a restaurant or food business. The project solves two connected needs. Customers need a simple way to browse products, create an account, complete a test payment, and review their orders. Business administrators need controlled tools for managing products, categories, orders, users, reports, and audit records. The technical team also needs repeatable deployment, security testing, monitoring, and a practical recovery process if the primary AWS Region becomes unavailable.")
    p(doc, "The solution was designed as a complete cloud system rather than a collection of unrelated pages. The browser communicates with managed AWS services, business logic runs in Lambda, data is stored in DynamoDB and S3, identities are handled by Cognito, payments are delegated to Stripe test mode, and operations are observed through CloudWatch. Terraform documents and recreates the infrastructure, while GitHub Actions checks code quality and security before any controlled deployment activity.")
    heading(doc, "1.2 Objectives", 2)
    for text in (
        "Provide a complete guest, customer, and administrator experience.",
        "Use managed and serverless AWS services to reduce server maintenance.",
        "Protect customer and administrator functions with authentication and authorization.",
        "Keep permanent AWS credentials and payment secrets out of source control.",
        "Validate application code and infrastructure automatically through GitHub Actions.",
        "Detect vulnerabilities, secrets, infrastructure misconfiguration, and code-quality problems.",
        "Create a documented cross-region recovery process with backup, restore, and validation evidence.",
        "Explain the design using the six AWS Well-Architected pillars."
    ):
        bullet(doc, text)
    heading(doc, "1.3 Users and main features", 2)
    add_table(doc, ["User", "Main capabilities"], [
        ("Guest", "Browse products and categories without signing in."),
        ("Registered customer", "Register, sign in, manage profile and addresses, use the cart, complete Stripe test checkout, and review orders."),
        ("Administrator/staff", "Use role-protected pages for menu, orders, users, reports, and audit records."),
        ("Technical operator", "Run CI/CD, monitor the platform, manage approved configuration, and perform recovery procedures."),
    ], [2100, 7260])
    heading(doc, "1.4 Scope and limitations", 2)
    p(doc, "The project demonstrates production-oriented architecture and security practices, but payments remain in Stripe test mode. The recovery exercise uses an isolated regional endpoint instead of changing public DNS. Cognito passwords, active sessions, refresh tokens, and MFA secrets cannot be exported, so recovered users must establish new credentials. SES sending permissions are regional and may remain restricted by the sandbox. These limitations are documented openly because a credible cloud design explains both capabilities and boundaries.")
    evidence(doc, "E01", "Project entry points", "the public eatinity.ca homepage and the GitHub repository README/structure", "proves that the application and organized source repository are available")

    heading(doc, "2. Solution Design", 1)
    heading(doc, "2.1 Functional requirements", 2)
    for text in (
        "Display products and categories to guests.", "Register, confirm, sign in, and sign out customers through Cognito.", "Maintain profile and address information.", "Add products and quantities to a client-side cart.", "Create a Stripe Checkout Session using server-authoritative product information.", "Receive and verify Stripe webhook events and update the order.", "Display customer order history.", "Restrict administrative functions using Cognito group claims.", "Manage products, categories, orders, users, reports, and audit events.", "Back up and restore DynamoDB and S3 data across Regions."
    ):
        bullet(doc, text)
    heading(doc, "2.2 Non-functional requirements", 2)
    for text in (
        "Use HTTPS, private frontend storage, and managed identity services.", "Apply least privilege to IAM roles and protected API routes.", "Use temporary OIDC credentials instead of long-lived AWS keys in GitHub.", "Store Stripe runtime secrets in Secrets Manager, outside source code and Terraform state.", "Use infrastructure as code, validation, and reviewed plans for repeatability.", "Enable encryption, versioning, point-in-time recovery, backup retention, and monitoring.", "Run linting, builds, unit tests, Trivy, SonarQube, and smoke tests.", "Target a four-hour recovery time objective and a 24-hour recovery point objective, subject to measured drill evidence."
    ):
        bullet(doc, text)
    heading(doc, "2.3 Important design decisions", 2)
    add_table(doc, ["Decision", "Reason"], [
        ("React and Vite", "A static single-page application supports both customer and administrator interfaces and can be delivered efficiently through a CDN."),
        ("S3, CloudFront, Route 53, and ACM instead of Amplify", "This preserves direct Terraform control over storage, caching, DNS, certificates, access, and recovery while supporting the eatinity.ca domain."),
        ("API Gateway and Lambda", "The API scales without permanently running servers and each function can receive focused IAM permissions."),
        ("DynamoDB", "A managed NoSQL database removes database-server administration and supports the project's access patterns."),
        ("Cognito JWT authorization", "Cognito manages identity; API Gateway validates tokens; group claims support administrator access control."),
        ("Stripe test mode", "Sensitive card processing remains with a specialized external payment provider while the project demonstrates checkout and webhook integration."),
        ("One GitHub repository and one main workflow", "The professor requires documented automation, not a specific number of repositories or workflow files. A monorepo keeps the submission understandable."),
        ("AWS OIDC", "Short-lived credentials remove the need to store AWS access keys in GitHub."),
        ("Cross-region AWS Backup", "Native recovery points and copies support an auditable regional recovery workflow."),
    ], [3000, 6360])

    heading(doc, "3. Architecture", 1)
    heading(doc, "3.1 End-to-end architecture", 2)
    p(doc, "A user enters eatinity.ca. Route 53 resolves the domain to CloudFront, and ACM provides the HTTPS certificate. CloudFront serves the React application from a private S3 website bucket through Origin Access Control. Product images are stored separately so website code and public business images can have different controls. AWS WAF is associated with CloudFront as a security control; it is not a DNS hop.")
    p(doc, "The React application calls API Gateway. Public product and category routes are available to guests. Protected customer and administrator routes use a Cognito JWT authorizer. API Gateway invokes Python Lambda functions, and Lambda execution roles permit only the required DynamoDB, S3, Secrets Manager, SNS, SES, and logging actions. DynamoDB stores products, categories, orders, users, and audit events.")
    evidence(doc, "E02", "Main AWS architecture", "docs/architecture/Architecture.png at full readable width", "explains all users, AWS services, security boundaries, databases, storage, and external services")
    heading(doc, "3.2 Customer and administrator request flows", 2)
    p(doc, "Guests can browse public information without a token. Registered customers sign in through Cognito and send a bearer JWT to protected API routes. Administrators use the same identity foundation, but the backend also checks authorized group claims. This layered design is important: the interface can hide restricted links, but the Lambda code still performs authorization because browser controls alone are not a security boundary.")
    heading(doc, "3.3 Payment and secret flow", 2)
    p(doc, "For checkout, the browser sends product identifiers and quantities. The checkout Lambda reloads the authoritative name, availability, and price from DynamoDB, calculates the order, and requests a Stripe Checkout Session. This prevents a customer from changing the price in browser data. Stripe later calls the public webhook endpoint. The webhook Lambda verifies the Stripe signature, handles duplicate events safely, updates the order, and triggers notifications.")
    p(doc, "The Stripe API key and webhook signing secret are stored together in AWS Secrets Manager. Terraform creates only the secret container; an authorized operator adds the value outside Terraform so the value does not enter state. Payment Lambdas receive the secret ARN and have GetSecretValue permission only for that resource. The frontend never receives the secret key.")
    heading(doc, "3.4 CI/CD and OIDC architecture", 2)
    p(doc, "The GitHub repository contains the frontend, backend, modular Terraform, recovery scripts, documentation, and workflows. Pushes and pull requests run validation and security gates. SonarQube uses a Windows self-hosted runner because the Community Edition server is local and cannot be reached from a GitHub-hosted runner. A separate manual OIDC verification workflow proves that GitHub can request a token and assume an AWS role without a stored access key.")
    evidence(doc, "E03", "CI/CD and OIDC architecture", "docs/architecture/CICD-OIDC-Architecture.png", "shows the relationship among developer, GitHub, security gates, protected Environment, OIDC, AWS, and post-deployment validation")
    heading(doc, "3.5 Disaster-recovery architecture", 2)
    p(doc, "The selected risk scenario is loss or serious unavailability of the primary us-east-1 Region. AWS Backup protects five DynamoDB tables and two S3 buckets. Recovery points are copied to a vault in ca-central-1. Terraform recreates regional application resources using Region-aware names. Data is restored into isolated drill targets, validated, and only then considered for controlled promotion. Cognito and SES require separate procedures because AWS Backup does not restore them in the same manner as DynamoDB and S3.")
    evidence(doc, "E04", "Disaster-recovery architecture", "docs/architecture/DisasterRecovery-Architecture.png", "shows primary protection, cross-region copy, isolated restore, validation, and service recovery")

    heading(doc, "4. Implementation", 1)
    heading(doc, "4.1 Repository and development environment", 2)
    p(doc, "The project uses one GitHub repository: MarjanHaghighi/eatinity. The root README explains each functional area. The frontend is under eatinity-frontend, the active and historical backend code is under eatinity-prod, modular infrastructure is under eatinity-iac, migration and recovery scripts are under eatinity-iac/migration, and evidence and instructions are under docs. Generated dependencies, Terraform state and plans, Lambda ZIP packages, build output, credentials, and secret files are excluded through .gitignore.")
    evidence(doc, "E05", "Repository organization", "the GitHub repository root with README and main folders visible", "proves that the source, infrastructure, recovery code, workflows, and documentation are organized in one submission repository")
    heading(doc, "4.2 Frontend validation", 2)
    p(doc, "The workflow installs exact frontend dependencies with npm ci, runs ESLint, and creates a production build with Vite. npm ci gives repeatable dependency installation from the lock file. Linting finds common code-quality problems before deployment, and the build confirms that the React application compiles into static assets.")
    heading(doc, "4.3 Backend validation and packaging", 2)
    p(doc, "Python unit tests validate important business rules such as checkout input, authoritative pricing, order-state transitions, administrator protection, audit behavior, reporting, webhook idempotency, and Secrets Manager integration. The workflow then creates a ZIP package for each Lambda directory while excluding Python cache files.")
    heading(doc, "4.4 Terraform validation", 2)
    p(doc, "Terraform formatting and validation run for the modular production environment and the separate GitHub OIDC bootstrap. Separating identity bootstrap from application deployment avoids a circular dependency: the OIDC role must exist before GitHub can use it to authenticate. A future deployment must initialize protected remote state, create a saved plan, review it, and apply that exact plan. For this submission, no new apply is required because the existing environments must not be changed.")
    evidence(doc, "E06", "Application and Terraform checks", "GitHub Actions run #14 with the expanded lint/build, 18 backend tests, Lambda packaging, Terraform formatting, and both validation steps", "proves that the application and infrastructure definitions passed automated validation")
    heading(doc, "4.5 Runtime operations", 2)
    p(doc, "CloudWatch receives Lambda logs and supports metrics and alarms. SNS provides operational alert delivery, and SES supports regional email identities and templates. Log retention is managed to avoid unlimited accumulation. Operators should correlate a test request with API Gateway/Lambda logs and verify that screenshots do not contain secrets or unnecessary customer information.")
    evidence(doc, "E07", "Operational monitoring", "CloudWatch log group with one successful request plus relevant metrics/alarms; crop account details if unnecessary", "proves observability and operational readiness")

    heading(doc, "5. Important Code and Configuration", 1)
    heading(doc, "5.1 Server-authoritative checkout", 2)
    p(doc, "The checkout Lambda does not trust a price received from the browser. It accepts identifiers and quantities, reads the current product record from DynamoDB, rejects unavailable or archived products, and uses the verified price when creating the Stripe session. This protects revenue and order integrity.")
    heading(doc, "5.2 Stripe webhook verification and idempotency", 2)
    p(doc, "The webhook route is public because Stripe must call it, but public does not mean unprotected. The Lambda validates the webhook signature with the secret from Secrets Manager. It also uses event and order state checks so a repeated webhook does not create duplicate processing.")
    heading(doc, "5.3 Administrator authorization", 2)
    p(doc, "The React interface includes protected routes and role-aware navigation for usability. The backend independently checks Cognito claims for every administrator Lambda. This prevents a customer from bypassing the interface and calling an administrator endpoint directly.")
    heading(doc, "5.4 Terraform safety and modularity", 2)
    p(doc, "The active modular Terraform separates identity, storage, database, application, delivery, operations, secrets, and recovery responsibilities. Production data controls include DynamoDB point-in-time recovery and deletion protection, S3 versioning and encryption, CloudFront private origin access, and narrowly scoped Lambda roles. Example variable files are committed; real variable, state, and plan files are ignored.")
    heading(doc, "5.5 GitHub Actions security gates", 2)
    p(doc, "The workflow grants read-only repository permission by default. Only the production deployment job requests id-token: write, and it can run only during a manual dispatch with deploy=true after the application, Trivy, and SonarQube jobs succeed. The separate OIDC verification workflow has no Terraform or deployment steps.")

    heading(doc, "6. User Walkthroughs", 1)
    heading(doc, "6.1 End-customer walkthrough", 2)
    p(doc, "The customer demonstration must be one connected journey. It should show the expected business result at each stage, not only a collection of page screenshots.")
    numbered_steps(doc, (
        "Register a test customer, confirm the account, and sign in.", "Browse categories and products, use search or filtering, and add items to the cart.", "Review quantities and totals, provide required checkout information, and start Stripe Checkout.", "Use Stripe test mode, return to the success page, and confirm that an order exists.", "Open order history, profile, and address pages.", "Sign out and verify that protected pages require authentication."
    ))
    evidence(doc, "E08", "Customer identity", "sign-up/confirmation or prepared customer sign-in and the authenticated account view", "proves customer authentication")
    evidence(doc, "E09", "Customer purchase journey", "products/search, cart, Stripe test checkout success, and created order; use two or three screenshots if needed", "proves a realistic connected customer workflow")
    evidence(doc, "E10", "Customer account results", "order history plus profile/address result", "proves that the customer can review persistent application data")
    heading(doc, "6.2 Client-administrator walkthrough", 2)
    numbered_steps(doc, (
        "Sign in with an administrator account and open the dashboard.", "Show that an ordinary customer is rejected from an administrator route.", "Create or update a category/product and verify the storefront result.", "Open an order, perform an allowed status transition, and explain invalid-transition protection.", "Review staff/users and super-admin protection.", "Open the sales report and select a period.", "Open the audit log and show the administrator action.", "Sign out and verify route protection."
    ))
    evidence(doc, "E11", "Administrator authorization", "administrator dashboard and an unauthorized-access rejection", "proves both successful authorization and denied access")
    evidence(doc, "E12", "Menu and order management", "product/category update and order workflow", "proves the main administrator business functions")
    doc.add_page_break()
    evidence(doc, "E13", "Users, reports, and audit", "user/staff page, sales report, and audit event", "proves governance and administrative visibility")

    heading(doc, "7. Disaster Recovery", 1)
    heading(doc, "7.1 Failure scenario, RPO, and RTO", 2)
    p(doc, "The selected scenario is failure of the primary AWS Region. The target recovery point objective is 24 hours because the scheduled backup plan runs daily. The target operator-led recovery time objective is four hours. These are design targets; the achieved values must be calculated from actual backup, copy, restore, deployment, and validation timestamps.")
    heading(doc, "7.2 Protection strategy", 2)
    p(doc, "AWS Backup creates native recovery points for five DynamoDB tables and two S3 buckets. Source recovery points are retained for 35 days and cross-region copies for 90 days by default. DynamoDB also uses point-in-time recovery, while S3 uses versioning and encryption. Backup jobs are monitored by ID, and local job records store non-secret job metadata rather than exported application data.")
    heading(doc, "7.3 Recovery process", 2)
    numbered_steps(doc, (
        "Confirm AWS account, source resources, destination Region, backup opt-in, versioning, and cost authorization.", "Identify or create source recovery points and verify that every job completed.", "Copy recovery points to the ca-central-1 recovery vault and verify completion.", "Restore DynamoDB into unique isolated drill tables and restore S3 into controlled versioned buckets.", "Compare table item counts, indexes, schema, object results, and business queries before promotion.", "Recreate regional application infrastructure from reviewed Terraform configuration.", "Synchronize transferable Cognito profiles and groups; require recovered users to establish new passwords.", "Recreate and validate regional SES identities, DKIM, templates, and notification topics.", "Deploy the regional frontend and run customer, administrator, API, and monitoring checks.", "Record actual RPO/RTO, limitations, final status, and evidence."
    ))
    heading(doc, "7.4 Why isolated restore is important", 2)
    p(doc, "A recovery point should not be connected directly to production traffic before validation. Restoring into isolated targets prevents damaged, incomplete, or incorrect data from replacing a working environment. Promotion is a separate decision after counts, indexes, application queries, identities, notifications, and business workflows are verified.")
    heading(doc, "7.5 Service-specific limitations", 2)
    p(doc, "Cognito passwords, sessions, refresh tokens, and MFA secrets are not exportable. The recovery process recreates the pool and groups, transfers allowed profile and membership information, and sends an approved reset flow. SES identity, production access, and quota state are regional; templates and configuration can be synchronized, but account-level approval may still be required. The drill does not change public DNS, which protects the working production site while recovery is being tested.")
    evidence(doc, "E14", "Backup and cross-region copy", "AWS Backup source/destination vaults and seven completed recovery points or copy jobs", "proves data was protected outside the primary Region")
    evidence(doc, "E15", "DynamoDB and S3 restore", "completed restore jobs plus table count/index validation and recovered S3 objects", "proves native restoration and data validation")
    evidence(doc, "E16", "Cognito and SES recovery", "Cognito users/groups validation and SES identity/DKIM/template status without personal data", "proves recovery of services not handled like DynamoDB/S3")
    evidence(doc, "E17", "Recovered application test", "recovery CloudFront URL, products API, customer/admin test, and CloudWatch result", "proves successful recovery rather than only a written plan")
    evidence(doc, "E18", "Measured recovery result", "a small table showing stage start/end times, achieved RTO, recovery-point time, achieved RPO, tester, and final status", "converts the recovery claim into measurable evidence")

    heading(doc, "8. Automated Security", 1)
    heading(doc, "8.1 Trivy", 2)
    p(doc, "Trivy runs on the complete repository for pushes, pull requests, and manual runs. It scans dependency vulnerabilities, embedded secrets, and Terraform/AWS misconfiguration. HIGH or CRITICAL findings produce exit code 1, so deployment cannot continue until the issue is fixed or a narrow, documented, expiring exception is approved.")
    p(doc, "During implementation, a React Router denial-of-service issue was remediated by upgrading react-router-dom and react-router from 7.17.0 to 7.18.2. Infrastructure improvements included SNS encryption and WAF attachment support. Remaining accepted findings are documented by scope, decision, and expiry instead of being silently ignored.")
    evidence(doc, "E19", "Trivy failure and finding", "an earlier failed workflow/log showing a real HIGH/CRITICAL finding", "demonstrates that the security gate can stop the pipeline")
    evidence(doc, "E20", "Trivy remediation and passing scan", "the remediation commit and expanded successful Trivy job from run #14", "proves the finding was addressed or formally reviewed and the gate passed")
    heading(doc, "8.2 SonarQube Community Edition", 2)
    p(doc, "SonarQube performs static analysis across React, Python, Terraform, and tests. A local Community Edition server was started in Docker. GitHub-hosted runners could not reach localhost, so a Windows self-hosted runner was registered. Early attempts exposed Windows action limitations, including missing PowerShell Core, GPG, and archive handling. The final workflow downloads and runs SonarScanner directly with Windows PowerShell, waits for the compute task, requests the Quality Gate result, and fails unless the status is OK.")
    p(doc, "The successful run proves that SonarQube is integrated into GitHub Actions rather than only executed manually. The dashboard showed a passed Quality Gate, while also reporting security, reliability, maintainability, duplication, and coverage information. Issues that are safe public API design decisions require review and documentation; real code defects should be fixed and rerun.")
    evidence(doc, "E21", "SonarQube GitHub Actions success", "successful run #14 with Run SonarQube scan and Enforce SonarQube Quality Gate expanded", "proves automated static analysis and gate enforcement")
    evidence(doc, "E22", "SonarQube dashboard", "docs/evidence/sonarcube/sonarqube-analysis-summary.png and quality-gate detail PNG", "proves the project result and Quality Gate status")
    evidence(doc, "E23", "SonarQube issue review/remediation", "one relevant issue detail plus the related code change or a documented safe-design review", "shows how analysis results were interpreted instead of only collecting a green badge")
    heading(doc, "8.3 Pass and fail behavior", 2)
    p(doc, "A green pipeline means the configured checks completed successfully; it does not mean the system has zero risk. A failed pipeline is useful evidence that a gate worked. The development process should show the original failure, the cause, the fix or documented decision, and the successful rerun. This creates an auditable improvement path.")

    heading(doc, "9. GitHub and AWS OIDC", 1)
    heading(doc, "9.1 Repository and workflow", 2)
    p(doc, "The main workflow runs application and Terraform checks, Trivy, SonarQube, and an optional production deployment. Validation runs on main pushes and pull requests. The SonarQube job is skipped for pull requests because the local self-hosted environment is configured for controlled project analysis. The deployment job requires a manual dispatch with deploy=true and depends on all required gates.")
    heading(doc, "9.2 OIDC authentication", 2)
    p(doc, "OIDC allows GitHub Actions to prove its identity to AWS without storing an AWS access key and secret key. GitHub requests a signed token for the sts.amazonaws.com audience. AWS checks the trusted provider and token claims, then AWS STS issues temporary role credentials. IAM permissions decide what that temporary identity is authorized to do.")
    p(doc, "The verified role is named eatinity-github-production-oidc. Its trust policy restricts the subject to the exact MarjanHaghighi/eatinity repository and production Environment. The verification role has zero attached and zero inline permission policies. This is intentional least privilege: it can prove authentication with sts:GetCallerIdentity but cannot change AWS resources. The verification workflow completed successfully and explicitly performed no Terraform plan, apply, or deployment.")
    evidence(doc, "E24", "Successful OIDC workflow", "the green AWS OIDC Identity Verification run with claims, credential configuration, GetCallerIdentity, and no-deployment confirmation", "proves temporary GitHub-to-AWS authentication")
    evidence(doc, "E25", "OIDC provider and trust policy", "IAM OIDC provider plus the role trust relationship showing audience and exact repository/production Environment subject", "proves that only the intended GitHub identity is trusted")
    evidence(doc, "E26", "Least-privilege verification role", "IAM role permissions showing zero attached policies and zero inline policies", "proves that the verification identity cannot modify AWS resources")
    heading(doc, "9.3 GitHub Environments and secrets", 2)
    p(doc, "The ci Environment stores the Sonar host URL as a variable and the Sonar token as a secret. The production Environment stores the OIDC verification role ARN as a non-secret variable. Environment configuration separates values by purpose and supports future approval rules. Screenshots must show names and protection settings but never reveal a secret value.")
    p(doc, "Runtime Stripe credentials do not belong in GitHub. They remain in Secrets Manager and are read at runtime by Lambda. This separates deployment identity, CI analysis credentials, and application runtime secrets. The repository and its history must be checked before every push, and exposed tokens must be revoked immediately.")
    evidence(doc, "E27", "GitHub Environments", "Settings > Environments for ci and production with variable/secret names only", "proves controlled configuration without exposing values")
    evidence(doc, "E28", "Secrets Manager metadata", "secret name/ARN, Region, rotation/configuration metadata, and last changed date only; never open Secret value", "proves secure runtime secret storage")

    heading(doc, "10. AWS Well-Architected Review", 1)
    p(doc, "The architecture was reviewed using the six AWS Well-Architected pillars. This review explains why the selected controls matter and where the project still has limits.")
    add_table(doc, ["Pillar", "How eatinity addresses it", "Remaining consideration"], [
        ("Operational Excellence", "Terraform, GitHub Actions, documented runbooks, CloudWatch, repeatable tests, audit records, and recovery procedures.", "Keep evidence current, measure recovery stages, and rehearse the operational runbook."),
        ("Security", "Cognito JWTs, backend role checks, least-privilege IAM, private S3 origin, HTTPS, WAF, Secrets Manager, OIDC, Trivy, SonarQube, and ignored secret/state files.", "Rotate exposed tokens, review accepted exceptions, improve test coverage, and keep screenshots free of secrets."),
        ("Reliability", "Serverless managed services, DynamoDB PITR/deletion protection, S3 versioning, AWS Backup, cross-region copies, isolated restore, and functional recovery testing.", "Measure achieved RPO/RTO and document DNS/failover decisions."),
        ("Performance Efficiency", "CloudFront caching, S3 static delivery, Lambda scaling, API Gateway, and DynamoDB managed capacity/indexes.", "Add repeatable load tests and tune Lambda memory, caching, and DynamoDB access patterns from measurements."),
        ("Cost Optimization", "Pay-per-use serverless services, static CDN delivery, managed encryption where appropriate, and controlled backup retention.", "Monitor backup, log, data-transfer, NAT-free, and self-hosted runner costs; remove unused drill resources through reviewed cleanup."),
        ("Sustainability", "Managed services, on-demand execution, CDN caching, minimized always-on compute, and controlled retention reduce unnecessary resource use.", "Shut down local SonarQube/runner when not needed and remove obsolete artifacts and recovery drills safely."),
    ], [1500, 5060, 2800])

    heading(doc, "11. Testing and Results", 1)
    p(doc, "The final evidence should record expected result, observed result, date, status, and evidence reference. The following results have already been demonstrated in the repository workflow and prior validation; live application and recovery results require the corresponding screenshots.")
    add_table(doc, ["Test area", "Current result", "Evidence to insert"], [
        ("Frontend lint/build", "Passed", "GitHub Actions run #14"),
        ("Backend unit tests", "18 tests passed", "Expanded GitHub Actions log"),
        ("Terraform formatting/validation", "Passed", "Expanded GitHub Actions log"),
        ("Trivy", "Passed after remediation/review", "Failed finding plus green rerun"),
        ("SonarQube", "Quality Gate passed in GitHub Actions", "Run #14 and dashboard PNGs"),
        ("OIDC", "Temporary identity verified; no AWS mutations", "Successful OIDC verification workflow, trust, and permissions"),
        ("Customer workflow", "[Insert observed result]", "E08-E10"),
        ("Administrator workflow", "[Insert observed result]", "E11-E13"),
        ("Disaster recovery", "[Insert measured result]", "E14-E18"),
        ("CloudWatch/operations", "[Insert observed result]", "E07"),
    ], [2800, 2800, 3760])

    heading(doc, "12. Risks, Limitations, and Future Improvements", 1)
    for text in (
        "The SonarQube server and self-hosted runner depend on a local computer; stop them when evidence collection is complete and protect the public-repository runner.",
        "SonarQube Community Build reports limited security analysis compared with commercial editions; Trivy and manual review remain complementary controls.",
        "Coverage is currently reported as 0% in SonarQube because coverage reports are not supplied; add JavaScript and Python coverage publishing.",
        "Cognito credentials and MFA cannot be transferred; communicate password reset steps during recovery.",
        "SES access and quotas are regional and may delay notification readiness.",
        "The DR drill avoids public DNS cutover; a future controlled exercise should validate documented DNS failover and rollback.",
        "A repeatable light load test should be added to measure response time, throttling, and scaling behavior.",
        "Accepted Trivy exceptions must be reviewed before their expiry dates and removed when a practical fix becomes available."
    ):
        bullet(doc, text)

    heading(doc, "13. Conclusion", 1)
    p(doc, "eatinity demonstrates a complete cloud solution that connects user experience, serverless application design, infrastructure as code, automated quality and security checks, temporary deployment identity, monitoring, and disaster recovery. The strongest design choice is separation of responsibilities: CloudFront and S3 deliver the interface, Cognito and API Gateway protect identity and access, Lambda applies business rules, DynamoDB and S3 store data, Secrets Manager protects runtime credentials, GitHub Actions validates changes, and AWS Backup supports regional recovery.")
    p(doc, "The project also documents the engineering process, including failed workflow attempts and their fixes. Trivy and SonarQube are not treated as screenshots only; they operate as pipeline gates. OIDC proves authentication without permanent AWS keys, and its zero-permission verification role demonstrates least privilege without risking either AWS Region. The final screenshots and recorded walkthrough will complete the professor's required pattern: explain the design, demonstrate that it works, and provide evidence.")

    heading(doc, "Appendix A. Evidence Checklist", 1)
    add_table(doc, ["ID", "Evidence", "Required?", "Inserted"], [
        ("E01-E07", "Repository, architecture, automated checks, and monitoring", "Yes", "[ ]"),
        ("E08-E13", "Customer and administrator walkthroughs", "Yes", "[ ]"),
        ("E14-E18", "Backup, restore, service recovery, application test, and measured RPO/RTO", "Yes", "[ ]"),
        ("E19-E23", "Trivy and SonarQube execution, results, and remediation/review", "Yes", "[ ]"),
        ("E24-E28", "OIDC, IAM, GitHub Environments, and Secrets Manager metadata", "Yes", "[ ]"),
        ("Supporting", "Additional Lambda/API/DynamoDB/CloudWatch screens", "Use when they prove an important claim", "[ ]"),
    ], [1100, 5000, 2100, 1160])
    show_box(doc, "Final privacy check", "Search the document for token patterns, AWS keys, passwords, secret values, personal customer data, and full payment details. Revoke any token that appeared in Remainder.docx or another shared artifact.", PALE_GOLD)

    heading(doc, "Appendix B. Repository Structure and References", 1)
    for text in (
        ".github/workflows/final-project-ci-cd.yml - application checks, Trivy, SonarQube, and optional deployment.",
        ".github/workflows/oidc-verification.yml - safe OIDC identity proof with no deployment.",
        "eatinity-frontend - React/Vite customer and administrator application.",
        "eatinity-prod - Python Lambda handlers and backend unit tests.",
        "eatinity-iac - modular Terraform, OIDC bootstrap, and migration/recovery scripts.",
        "docs/architecture - production, CI/CD/OIDC, and disaster-recovery diagrams.",
        "docs/evidence - sanitized submission evidence.",
        "sonar-project.properties - SonarQube analysis scope and exclusions.",
        "docs/SECURITY_EXCEPTIONS.md - narrow, reviewed, expiring Trivy exceptions.",
        "Repository: https://github.com/MarjanHaghighi/eatinity",
        "Application: https://eatinity.ca",
        "Video: [Insert tested unlisted YouTube URL]"
    ):
        bullet(doc, text)

    path = OUT / "eatinity_Final_Project_Report_Draft.docx"
    doc.save(path)
    return path


def video_document():
    doc = Document()
    configure(doc, "eatinity | 15-Minute Video Demonstration Script")
    cover(doc, "15-Minute Video Demonstration Script", "Exact speaking guidance, screen-sharing cues, timing, and evidence checklist")

    heading(doc, "How to Use This Script", 1)
    p(doc, "Open every page and console tab before recording. Use a prepared customer account, administrator account, Stripe test session, successful GitHub Actions runs, SonarQube dashboard, OIDC evidence, and DR evidence. Keep the self-hosted runner and SonarQube running only when needed. Hide bookmarks, notifications, account numbers, email addresses, tokens, secret values, and unrelated browser tabs.")
    p(doc, "The script totals approximately 15 minutes. Speak naturally and do not read file paths character by character. If a live action is slow, show a prepared successful result and explain what occurred. Every group member must appear on camera and speak. Assign the sections below before recording.")
    show_box(doc, "Before recording", "Test microphone, camera, screen sharing, application login, Stripe test mode, GitHub links, SonarQube localhost access, AWS console session, recovery evidence, and the unlisted upload process. Close Remainder.docx because it contains a token that must be revoked and removed.", PALE_GOLD)

    heading(doc, "15-Minute Run of Show", 1)
    add_table(doc, ["Time", "Section", "Screen"], [
        ("0:00-1:15", "Problem, objectives, and result", "Title slide + eatinity.ca"),
        ("1:15-3:15", "Architecture and six pillars", "Main architecture diagram"),
        ("3:15-6:00", "Customer journey", "Live website + Stripe test mode"),
        ("6:00-8:15", "Administrator journey", "Admin pages + authorization"),
        ("8:15-10:30", "Disaster recovery", "DR diagram + AWS Backup/restore evidence"),
        ("10:30-13:45", "GitHub Actions, Trivy, SonarQube, OIDC, secrets", "GitHub, SonarQube, IAM/Environment"),
        ("13:45-15:00", "Results, limitations, conclusion", "Results summary + repository"),
    ], [1400, 3300, 4660])

    heading(doc, "0:00-1:15 | Opening: Problem, Objectives, and Solution", 1)
    show_box(doc, "SHOW", "Start with a simple title slide, then open https://eatinity.ca.")
    p(doc, "Hello Professor. We are presenting eatinity, our secure serverless e-commerce platform built on AWS. The project supports three main experiences: guests can browse products, registered customers can complete a full test purchase and manage their account, and authorized staff can manage the business through a protected administrator portal.")
    p(doc, "Our objective was not only to build application pages. We wanted a complete cloud solution with secure identity, trusted payments, repeatable infrastructure, automated testing, monitoring, and disaster recovery. The application uses managed AWS services so that the team can focus on business logic instead of maintaining servers.")
    p(doc, "During this demonstration, we will explain the architecture, complete customer and administrator workflows, show cross-region recovery evidence, and demonstrate GitHub Actions with Trivy, SonarQube, and AWS OIDC.")
    show_box(doc, "ACTION", "Point briefly to the working homepage. Do not spend time browsing yet.", GREEN)

    heading(doc, "1:15-3:15 | Architecture and Design Decisions", 1)
    show_box(doc, "SHOW", "Open docs/architecture/Architecture.png at a readable zoom.")
    p(doc, "A user enters eatinity.ca. Route 53 directs the domain to CloudFront, and AWS Certificate Manager provides HTTPS. CloudFront serves our React application from a private S3 bucket through Origin Access Control. We use S3 and CloudFront instead of Amplify because we already own the domain and wanted direct Terraform control over hosting, caching, certificates, security, and recovery.")
    p(doc, "The browser calls API Gateway. Public product routes support guests. Protected routes use Cognito JWT tokens. API Gateway invokes Python Lambda functions, and DynamoDB stores products, categories, orders, users, and audit records. The backend checks administrator group claims; hiding a menu item in React is not considered security.")
    p(doc, "For checkout, Lambda reads the real product price from DynamoDB instead of trusting browser data. It creates a Stripe test Checkout Session. Stripe later sends a signed webhook to API Gateway, and the webhook Lambda verifies the signature before updating the order. Stripe secrets are stored in Secrets Manager and never sent to the frontend.")
    p(doc, "This design supports all six AWS Well-Architected pillars. Automation and monitoring support operational excellence. Cognito, IAM, OIDC, WAF, Secrets Manager, Trivy, and SonarQube support security. Backup and cross-region recovery support reliability. CloudFront and serverless scaling support performance. Pay-per-use managed services and controlled retention support cost optimization. Managed services and reduced always-on compute support sustainability.")
    show_box(doc, "ACTION", "Follow only the main request path with your cursor. Do not explain every small icon.", GREEN)

    heading(doc, "3:15-6:00 | End-Customer Demonstration", 1)
    show_box(doc, "SHOW", "Live eatinity.ca customer workflow. Use a prepared account to avoid waiting for email.")
    p(doc, "Now I will demonstrate the complete customer journey. I start by signing in with a prepared test customer. Authentication is handled by Amazon Cognito. After sign-in, the application keeps the session token and sends it only to protected API routes.")
    show_box(doc, "ACTION", "Sign in, show the account name briefly, then return to products.", GREEN)
    p(doc, "The customer can browse categories, search or filter products, and add products to the cart. The cart is managed in the customer interface, but final pricing is not trusted from the browser.")
    show_box(doc, "ACTION", "Search for one item, add two items, open the cart, change one quantity, and show the total.", GREEN)
    p(doc, "When I continue to checkout, the backend reloads each product from DynamoDB and calculates the trusted total. It then creates a Stripe Checkout Session using the secret key from Secrets Manager. The customer enters payment information only on Stripe's test page, so eatinity does not handle card data directly.")
    show_box(doc, "ACTION", "Open Stripe Checkout and use the prepared Stripe test card. Never show a real card or personal address.", GREEN)
    p(doc, "After the test payment, Stripe sends a signed webhook. Our webhook verifies the signature, prevents duplicate processing, updates the order, and supports notifications. The customer returns to the success page and can see the order in order history. Profile and address pages show the remaining account functions.")
    show_box(doc, "ACTION", "Show success, order history, profile/address, then sign out. If payment is slow, use prepared success and order-history tabs.", GREEN)

    heading(doc, "6:00-8:15 | Client-Administrator Demonstration", 1)
    show_box(doc, "SHOW", "Open the administrator sign-in and dashboard with prepared data.")
    p(doc, "The administrator portal uses the same Cognito identity service but requires approved group claims. The React route improves the user experience, and each administrator Lambda also performs authorization on the backend. This prevents a customer from calling a protected administrator API directly.")
    show_box(doc, "ACTION", "Show a prepared unauthorized rejection first, then sign in as an administrator.", GREEN)
    p(doc, "From the dashboard, an administrator can manage categories and products. I will update one product and confirm that the change appears in the storefront. Product deletion is designed as archive behavior where appropriate so that business history is not lost.")
    show_box(doc, "ACTION", "Update one safe field, save it, and show the result. Use prepared evidence if you do not want to change production data during recording.", GREEN)
    p(doc, "The Orders page applies a controlled status workflow. Allowed transitions can proceed, while invalid transitions are rejected. The Users page supports staff management and protects the super-admin account. Reports summarize paid sales for the selected period, and the audit log records important administrator actions.")
    doc.add_page_break()
    show_box(doc, "ACTION", "Show one order, users/staff, report, and the related audit entry. Then sign out.", GREEN)

    heading(doc, "8:15-10:30 | Disaster-Recovery Demonstration", 1)
    show_box(doc, "SHOW", "Open docs/architecture/DisasterRecovery-Architecture.png, then prepared AWS Backup and restore evidence.")
    p(doc, "Our selected failure scenario is loss of the primary us-east-1 Region. Our target recovery point objective is 24 hours because backup is scheduled daily, and our target operator-led recovery time objective is four hours. The achieved values must come from actual job and validation timestamps.")
    p(doc, "AWS Backup protects five DynamoDB tables and two S3 buckets. Recovery points are copied to a vault in ca-central-1. We first restore into isolated drill resources. This is important because we do not connect unvalidated recovered data to users. We compare table counts, indexes, schema, S3 objects, and business queries before any promotion decision.")
    show_box(doc, "ACTION", "Show completed source/copy recovery points, one DynamoDB restore, item/index validation, and recovered S3 objects.", GREEN)
    p(doc, "Terraform recreates the regional application components without changing the working primary environment. Cognito requires a separate process because passwords, sessions, refresh tokens, and MFA secrets cannot be exported. We recreate users and groups from allowed profile data and require a password reset. SES is also regional, so identities, DKIM, templates, and account status must be validated in the recovery Region.")
    show_box(doc, "ACTION", "Show Cognito group/user validation and SES/DKIM/template evidence, with emails and account details cropped.", GREEN)
    p(doc, "Finally, we test the recovered CloudFront URL, products API, customer and administrator functions, and CloudWatch results. We do not change public DNS during the drill. Our evidence records the recovery point, stage times, final validation, limitations, and lessons learned.")
    show_box(doc, "ACTION", "Show the recovered application/API result and a small measured RPO/RTO summary.", GREEN)

    heading(doc, "10:30-13:45 | GitHub Actions, Security, OIDC, and Secrets", 1)
    show_box(doc, "SHOW", "Open successful GitHub Actions run #14.")
    p(doc, "We use one GitHub repository and one main workflow because the assignment requires clear automation, not a specific number of repositories. On pushes and pull requests, the first job installs dependencies, runs frontend lint and build, executes 18 backend unit tests, packages Lambda functions, and validates both production and OIDC Terraform.")
    show_box(doc, "ACTION", "Expand the successful application and Terraform steps. Keep the green status and step names visible.", GREEN)
    p(doc, "The Trivy job scans the entire repository for dependency vulnerabilities, embedded secrets, and infrastructure misconfiguration. HIGH and CRITICAL findings return a failure code and block deployment. We experienced a real dependency finding, upgraded React Router from 7.17.0 to 7.18.2, documented narrow expiring exceptions, and obtained a successful final scan.")
    show_box(doc, "ACTION", "Show the earlier failed finding, remediation commit, and expanded green Trivy job.", GREEN)
    p(doc, "SonarQube performs static analysis for code quality and security. Our Community Edition server runs locally in Docker. A GitHub-hosted runner cannot reach localhost, so we configured a Windows self-hosted runner. We also resolved Windows scanner problems and finally ran SonarScanner directly. The workflow waits for the server analysis and fails unless the Quality Gate is OK.")
    show_box(doc, "ACTION", "Expand Run SonarQube scan and Enforce SonarQube Quality Gate, then show the dashboard with Passed.", GREEN)
    p(doc, "For AWS authentication, GitHub Actions uses OIDC instead of permanent AWS keys. GitHub requests a signed identity token, AWS checks its claims, and STS issues temporary credentials. Our trust policy is restricted to the exact eatinity repository and production Environment.")
    show_box(doc, "ACTION", "Open the successful OIDC verification workflow, IAM trust relationship, and GitHub production Environment.", GREEN)
    p(doc, "The verification role has zero attached and zero inline permission policies. This is deliberate least privilege: it proves the identity connection with GetCallerIdentity but cannot change AWS resources. The workflow confirms that no Terraform plan, apply, or deployment occurred. Runtime Stripe secrets are separate and remain in Secrets Manager, where only the payment Lambdas can read the value.")
    show_box(doc, "ACTION", "Show zero IAM policies and Secrets Manager metadata only. Never click Reveal secret value.", GREEN)

    heading(doc, "13:45-15:00 | Results, Limitations, and Conclusion", 1)
    show_box(doc, "SHOW", "Return to a simple results slide or the report testing table, then the GitHub repository.")
    p(doc, "Our automated results show that the frontend lint and build passed, all 18 backend tests passed, Terraform validation passed, Trivy passed after remediation and review, SonarQube completed inside GitHub Actions with a passed Quality Gate, and AWS OIDC successfully issued a temporary verification identity without resource permissions.")
    p(doc, "The main limitations are that SonarQube and the self-hosted runner depend on a local computer, SonarQube coverage reports are not yet supplied, Cognito passwords and MFA cannot be transferred, SES approval is regional, and our DR drill does not change public DNS. These are documented so another technical team understands what is proven and what still requires operational planning.")
    p(doc, "In conclusion, eatinity connects a working customer and administrator application with secure AWS architecture, infrastructure as code, automated security gates, temporary cloud identity, monitoring, and cross-region recovery. Our project follows the professor's required approach: we explain each design decision, demonstrate the working process, and provide evidence. Thank you.")
    show_box(doc, "ACTION", "End on the repository URL and state that the YouTube link was tested in an Incognito window.", GREEN)

    heading(doc, "Screen Preparation Checklist", 1)
    for text in (
        "Title slide with project, group, course, and repository URL.", "Main, CI/CD/OIDC, and DR architecture diagrams open at readable zoom.", "Prepared customer and administrator test accounts.", "Products, cart, Stripe test checkout, order history, profile, admin pages, reports, and audit data.", "Successful GitHub Actions run #14 and earlier failed/remediated security evidence.", "SonarQube Passed dashboard and Quality Gate details.", "Successful OIDC verification run, trust policy, zero permissions, and GitHub Environments.", "Secrets Manager metadata only.", "AWS Backup/copy/restore, DynamoDB/S3 validation, Cognito, SES, recovered application, and RPO/RTO evidence.", "All sensitive values cropped; Remainder.docx token revoked.", "All members assigned speaking sections, cameras on, and screen sharing rehearsed.", "Final YouTube upload set to Unlisted and tested in an Incognito window."
    ):
        bullet(doc, text)

    heading(doc, "Recommended Speaker Assignment", 1)
    add_table(doc, ["Speaker", "Suggested section", "Time"], [
        ("Member 1", "Opening and architecture", "0:00-3:15"),
        ("Member 2", "Customer and administrator demonstrations", "3:15-8:15"),
        ("Member 3", "Disaster recovery", "8:15-10:30"),
        ("Member 4 or shared", "CI/CD, security, OIDC, and conclusion", "10:30-15:00"),
    ], [1900, 5660, 1800])
    p(doc, "If the group has a different number of members, divide the sections so every member appears on camera and speaks. Keep transitions short: the next speaker should already have the correct tab ready.")

    doc.add_page_break()
    heading(doc, "Emergency Backup Plan During Recording", 1)
    for text in (
        "If eatinity.ca is slow, show a prepared successful customer-flow recording or screenshots and explain the expected result.",
        "If Stripe is slow, show the prepared Checkout page, success result, webhook/order evidence, and clearly state that it is test mode.",
        "If SonarQube or the runner is unavailable, show the saved green GitHub Actions run and sanitized dashboard PNGs.",
        "If AWS Backup pages are slow, show timestamped prepared evidence and the job IDs/status records without secret data.",
        "Do not troubleshoot for several minutes during the final recording. State the evidence, continue the story, and keep within 15 minutes."
    ):
        bullet(doc, text)

    path = OUT / "eatinity_15-Minute_Video_Demo_Script.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(report_document())
    print(video_document())
